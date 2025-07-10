from collections.abc import Generator
from typing import Any
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from utils.digitforce_api import DigitForceApi
import base64
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
import re
MAX_QUERY_LENGTH = 1000
API_MODULE_NAME = "PythonDataAnalysis"

class DataAnalysisTool(Tool):
    def _process_markdown_to_word(self, doc: Document, text: str):
        # Split text into lines
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Handle headers
            if line.startswith('#'):
                level = len(re.match('^#+', line).group())
                text = line.lstrip('#').strip()
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(16 - level)  # h1=15pt, h2=14pt, etc.
                run.font.bold = True
                i += 1
                continue

            # Handle tables
            if line.startswith('|'):
                # Collect all table lines
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) >= 3:  # Header, separator, and at least one row
                    # Create table
                    num_cols = len(table_lines[0].split('|')) - 2  # -2 because of leading and trailing |
                    table = doc.add_table(rows=len(table_lines)-1, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    # Process header
                    header_cells = table_lines[0].split('|')[1:-1]
                    for j, cell in enumerate(header_cells):
                        self._add_markdown_bold_to_run(table.cell(0, j).paragraphs[0], cell.strip())
                    
                    # Process data rows (skip separator line)
                    for row_idx, row_line in enumerate(table_lines[2:], 1):
                        cells = row_line.split('|')[1:-1]
                        for col_idx, cell in enumerate(cells):
                            self._add_markdown_bold_to_run(table.cell(row_idx, col_idx).paragraphs[0], cell.strip())
                continue

            # Handle bold text
            if '**' in line or '__' in line:
                paragraph = doc.add_paragraph()
                self._add_markdown_bold_to_run(paragraph, line)
                i += 1
                continue

            # Handle regular text
            if line:
                paragraph = doc.add_paragraph()
                self._add_markdown_bold_to_run(paragraph, line)
            else:
                doc.add_paragraph()  # Add empty line for spacing
            i += 1

    def _add_markdown_bold_to_run(self, paragraph, text):
        # 支持 **加粗** 和 __加粗__
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('__') and part.endswith('__'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = paragraph.add_run(part)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        try:
            # Retrieve the API key
            api_key = self.runtime.credentials["digitforce_api_key"]
            # Extract variables from the tool_parameters
            query = tool_parameters['query']
            input_data = tool_parameters.get('input_data')
            file = tool_parameters.get('input_file')

            if not input_data and not file:
                yield self.create_text_message("Please provide input_data or file")
                return

            if len(query) > MAX_QUERY_LENGTH:
                yield self.create_text_message("The query is too long. Please check if you have entered any incorrect variables")
                return

            post_data = {
                "query": query,
                "input_data": input_data
            }
            if file:
                post_data['file_extension'] = file.extension.lower()
                post_data['file_url'] = file.url
                if file.blob:
                    post_data['file_blob'] = base64.b64encode(file.blob).decode('utf-8')              
                post_data['sheet_name'] = None  # None means read all sheets
            result = DigitForceApi(api_key).dify_api_post(post_data, service_name=API_MODULE_NAME)
            if result:
                yield self.create_text_message(result)
                try:
                    # Create a new Word document
                    doc = Document()
                    
                    # Set default font for the document
                    style = doc.styles['Normal']
                    style.font.name = '微软雅黑'
                    style.font.size = Pt(12)
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    
                    # Process the markdown content
                    self._process_markdown_to_word(doc, result)
                    
                    # Save the document to a BytesIO object
                    docx_bytes = BytesIO()
                    doc.save(docx_bytes)
                    docx_bytes.seek(0)
                    
                    # Yield the document as a blob message
                    yield self.create_blob_message(
                        blob=docx_bytes.getvalue(),
                        meta={
                            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            "file_name": "analysis_result.docx"
                        }
                    )
                    docx_bytes.close()
                except:
                    yield self.create_text_message("word download error")
            else:
                yield self.create_text_message("Result is empty.Possible data source or code generation issue. Check input data and retry.")
        except Exception as e:
            yield self.create_text_message(str(e))

